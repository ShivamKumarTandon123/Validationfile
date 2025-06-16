from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename
from docx import Document
import os
import zipfile
import xml.etree.ElementTree as ET

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Check font and size in normal paragraphs
def check_fonts(doc):
    font_name_issues = []
    font_size_issues = []

    for i, para in enumerate(doc.paragraphs):
        para_text = para.text.strip()
        found_name_issue = False
        found_size_issue = False

        for run in para.runs:
            font = run.font
            # Font name check
            if not found_name_issue and font.name and font.name != "Times New Roman":
                font_name_issues.append(
                    f"Paragraph {i+1}: Font '{font.name}' instead of 'Times New Roman'. Text: \"{para_text}\""
                )
                found_name_issue = True

            # Font size check
            if not found_size_issue and font.size and font.size.pt != 12:
                font_size_issues.append(
                    f"Paragraph {i+1}: Font size {font.size.pt}pt instead of 12pt. Text: \"{para_text}\""
                )
                found_size_issue = True

            if found_name_issue and found_size_issue:
                break  # no need to keep checking this paragraph

    return font_name_issues, font_size_issues

# Check font size in tables
def check_table_fonts(doc):
    warnings = []
    for t_index, table in enumerate(doc.tables):
        issue_found = False
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        size = run.font.size
                        if size and size.pt != 9:
                            warnings.append(f"Table {t_index+1}: Found font size {size.pt}pt instead of 9pt")
                            issue_found = True
                            break
                    if issue_found:
                        break
                if issue_found:
                    break
            if issue_found:
                break
    return warnings

# Check page orientation
def check_orientation(doc):
    try:
        for section in doc.sections:
            if "PORTRAIT" not in section.orientation.__str__():
                return ["Page orientation is not portrait."]
    except:
        return ["Unable to determine document orientation."]
    return []

# Check TOC presence and hyperlink anchors
def check_toc_links(docx_path):
    toc_links = []
    bookmarks = set()
    missing_links = []

    try:
        with zipfile.ZipFile(docx_path, 'r') as docx:
            xml_content = docx.read("word/document.xml")
            root = ET.fromstring(xml_content)
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            # Collect all hyperlink anchors in TOC
            for link in root.findall('.//w:hyperlink', ns):
                anchor = link.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}anchor')
                if anchor and anchor.startswith('_Toc'):
                    texts = [t.text for t in link.findall('.//w:t', ns) if t.text]
                    visible_text = ' '.join(texts)
                    toc_links.append(anchor)

            # Collect all bookmarks (targets)
            for bookmark in root.findall('.//w:bookmarkStart', ns):
                name = bookmark.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name')
                if name:
                    bookmarks.add(name)

            # Check which anchors don’t match bookmarks
            for anchor in toc_links:
                if anchor not in bookmarks:
                    missing_links.append(anchor)

        return {
            "toc_present": bool(toc_links),
            "total_links_in_toc": len(toc_links),
            "valid_links": len(toc_links) - len(missing_links),
            "broken_links": len(missing_links),
            "broken_anchor_ids": missing_links
        }

    except Exception as e:
        return {"error": f"TOC parsing error: {str(e)}"}

def check_margins(doc):
    errors = []
    warnings = []

    for i, section in enumerate(doc.sections):
        margins = {
            "top": section.top_margin.inches,
            "bottom": section.bottom_margin.inches,
            "left": section.left_margin.inches,
            "right": section.right_margin.inches,
        }

        for side, value in margins.items():
            if value < 0.75:
                errors.append(f"Section {i+1}: {side.capitalize()} margin is {value:.2f}in (less than 0.75in)")
            elif value != 1.0:
                warnings.append(f"Section {i+1}: {side.capitalize()} margin is {value:.2f}in (not exactly 1in)")

    return errors, warnings

def check_header_footer_distance(doc):
    errors = []

    for i, section in enumerate(doc.sections):
        if section.header_distance.inches < 0.38:
            errors.append(f"Section {i+1}: Header is {section.header_distance.inches:.2f}in from top (must be ≥ 0.38in)")
        if section.footer_distance.inches < 0.38:
            errors.append(f"Section {i+1}: Footer is {section.footer_distance.inches:.2f}in from bottom (must be ≥ 0.38in)")

    return errors

def check_internal_links(docx_path):
    internal_links = []
    doc = Document(docx_path)
    all_text = " ".join(p.text for p in doc.paragraphs).lower()

    with zipfile.ZipFile(docx_path, 'r') as docx:
        xml = docx.read("word/document.xml")
        root = ET.fromstring(xml)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        for link in root.findall('.//w:hyperlink', ns):
            anchor = link.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}anchor')
            if anchor:
                link_text = ''.join(t.text for t in link.findall('.//w:t', ns) if t.text).strip()
                if anchor.startswith('_Toc'):
                    # TOC links are handled separately
                    continue
                if link_text and link_text.lower() not in all_text:
                    internal_links.append(f"Broken link: \"{link_text}\" does not appear anywhere in document")

    return internal_links

def check_external_links(docx_path, timeout=5):
    invalid_links = []

    with zipfile.ZipFile(docx_path, 'r') as docx:
        xml = docx.read("word/document.xml")
        root = ET.fromstring(xml)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        for link in root.findall('.//w:hyperlink', ns):
            r_id = link.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            if r_id:
                # Follow relationship to get actual hyperlink
                try:
                    rels = docx.read("word/_rels/document.xml.rels")
                    rel_root = ET.fromstring(rels)
                    rel_ns = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}

                    for rel in rel_root.findall('r:Relationship', rel_ns):
                        if rel.attrib.get('Id') == r_id:
                            target = rel.attrib.get('Target')
                            if target and target.startswith('http'):
                                try:
                                    response = requests.head(target, timeout=timeout, allow_redirects=True)
                                    if response.status_code >= 400:
                                        invalid_links.append(f"Broken link: {target} (HTTP {response.status_code})")
                                except Exception as e:
                                    invalid_links.append(f"Invalid link: {target} ({e})")
                except Exception as e:
                    invalid_links.append(f"Could not resolve external link relationships: {e}")

    return invalid_links

# Upload form route
@app.route('/')
def index():
    return render_template('upload.html')

# Upload and process file
@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    if not file.filename.endswith('.docx'):
        return jsonify({"error": "Only .docx files are supported"}), 400

    filename = secure_filename(file.filename)
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(path)

    try:
        doc = Document(path)
        font_name_issues, font_size_issues = check_fonts(doc)
        margin_errors, margin_warnings = check_margins(doc)
        header_footer_errors = check_header_footer_distance(doc)
        internal_link_issues = check_internal_links(path)
        external_link_issues = check_external_links(path)
        result = {
            "font_name_issues": font_name_issues,
            "font_size_issues": font_size_issues,
            "table_font_issues": check_table_fonts(doc),
            "orientation_issues": check_orientation(doc),
            "toc_check": check_toc_links(path),
            "margin_errors": margin_errors,
            "margin_warnings": margin_warnings,
            "header_footer_errors": header_footer_errors,
            "internal_link_issues": internal_link_issues,
            "external_link_issues": external_link_issues,
        }
        return render_template("result.html", result=result, filename=filename)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

def test():
    filename = 'data_text.docx'
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    print(path)
    try:
        doc = Document(path)
        font_name_issues, font_size_issues = check_fonts(doc)
        margin_errors, margin_warnings = check_margins(doc)
        header_footer_errors = check_header_footer_distance(doc)
        internal_link_issues = check_internal_links(path)
        external_link_issues = check_external_links(path)
        result = {
            "font_name_issues": font_name_issues,
            "font_size_issues": font_size_issues,
            "table_font_issues": check_table_fonts(doc),
            "orientation_issues": check_orientation(doc),
            "toc_check": check_toc_links(path),
            "margin_errors": margin_errors,
            "margin_warnings": margin_warnings,
            "header_footer_errors": header_footer_errors,
            "internal_link_issues": internal_link_issues,
            "external_link_issues": external_link_issues,
        }
        print(result)
        print(result)
    except Exception as e:
        print(e)

if __name__ == '__main__':
    app.run(
        host='0.0.0.0',
        debug=True,
        port=5001,
    )
    # test()
